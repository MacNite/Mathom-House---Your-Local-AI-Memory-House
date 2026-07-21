"""Unit tests for local document text extraction.

These run with no network and no models: they exercise the extraction and the
zip/PDF safety guards directly. The ``client`` fixture is used only to ensure
settings (limits, data dir) are initialised the same way the app sees them.
"""

import zipfile
from pathlib import Path

import pytest
from fastapi.testclient import TestClient

from app.services import documents


def test_extract_txt(tmp_path: Path, client: TestClient) -> None:
    path = tmp_path / "note.txt"
    path.write_text("Hello archive.\nSecond line.", encoding="utf-8")
    assert documents.extract_text(path, ".txt") == "Hello archive.\nSecond line."


def test_extract_markdown_with_bom(tmp_path: Path, client: TestClient) -> None:
    path = tmp_path / "note.md"
    # Leading UTF-8 BOM must be stripped by utf-8-sig decoding.
    path.write_bytes(b"\xef\xbb\xbf# Title\n\nBody")
    assert documents.extract_text(path, ".md") == "# Title\n\nBody"


def test_text_rejects_nul_bytes(tmp_path: Path, client: TestClient) -> None:
    path = tmp_path / "binary.txt"
    path.write_bytes(b"text\x00more")
    with pytest.raises(documents.DocumentExtractionError):
        documents.extract_text(path, ".txt")


def test_text_rejects_empty(tmp_path: Path, client: TestClient) -> None:
    path = tmp_path / "blank.txt"
    path.write_text("   \n\t", encoding="utf-8")
    with pytest.raises(documents.DocumentExtractionError):
        documents.extract_text(path, ".txt")


def test_text_rejects_over_limit(
    tmp_path: Path, client: TestClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    from app.config import get_settings

    monkeypatch.setattr(get_settings(), "max_text_chars", 10)
    path = tmp_path / "long.txt"
    path.write_text("x" * 50, encoding="utf-8")
    with pytest.raises(documents.DocumentExtractionError):
        documents.extract_text(path, ".txt")


def test_unsupported_extension(tmp_path: Path, client: TestClient) -> None:
    path = tmp_path / "thing.rtf"
    path.write_text("data", encoding="utf-8")
    with pytest.raises(documents.DocumentExtractionError):
        documents.extract_text(path, ".rtf")


def test_docx_roundtrip(tmp_path: Path, client: TestClient) -> None:
    from docx import Document

    document = Document()
    document.add_paragraph("First paragraph.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "left"
    table.cell(0, 1).text = "right"
    path = tmp_path / "doc.docx"
    document.save(path)

    extracted = documents.extract_text(path, ".docx")
    assert "First paragraph." in extracted
    assert "left\tright" in extracted


def test_docx_rejects_non_zip(tmp_path: Path, client: TestClient) -> None:
    path = tmp_path / "fake.docx"
    path.write_bytes(b"not a zip at all")
    with pytest.raises(documents.DocumentExtractionError):
        documents.extract_text(path, ".docx")


def test_docx_rejects_zip_without_document(tmp_path: Path, client: TestClient) -> None:
    path = tmp_path / "wrong.docx"
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("hello.txt", "not really a docx")
    with pytest.raises(documents.DocumentExtractionError):
        documents.extract_text(path, ".docx")


def test_docx_rejects_too_many_entries(
    tmp_path: Path, client: TestClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(documents, "MAX_ZIP_ENTRIES", 2)
    path = tmp_path / "bomb.docx"
    with zipfile.ZipFile(path, "w") as archive:
        for i in range(5):
            archive.writestr(f"file-{i}.xml", "x")
    with pytest.raises(documents.DocumentExtractionError):
        documents.extract_text(path, ".docx")


def test_pdf_rejects_invalid_bytes(tmp_path: Path, client: TestClient) -> None:
    path = tmp_path / "broken.pdf"
    path.write_bytes(b"%PDF-1.4 but not really")
    with pytest.raises(documents.DocumentExtractionError):
        documents.extract_text(path, ".pdf")
